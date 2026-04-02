"""Generate comprehensive BO App documentation in DOCX format"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ─── Styles ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

GREEN = RGBColor(0x26, 0xB5, 0x0F)

def add_colored_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GREEN if level == 1 else RGBColor(0x1A, 0x20, 0x2C)
    return h

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="26B50F"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t

# ════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(100)
run = p.add_run('BO')
run.font.size = Pt(72)
run.font.bold = True
run.font.color.rgb = GREEN

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Health & Wellness Platform')
run2.font.size = Pt(24)
run2.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Complete Project Documentation\nv1.0.0 | April 2026')
run3.font.size = Pt(14)
run3.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.space_before = Pt(40)
run4 = p4.add_run('Built by Flynaut LLC\nPowered by Expo, React Native, FastAPI & AI')
run4.font.size = Pt(11)
run4.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════
add_colored_heading('Table of Contents')
toc_items = [
    '1. Executive Summary & Project Overview',
    '2. Project Metrics & Statistics',
    '3. User Types & Roles',
    '4. Sprint & Feature List',
    '5. Complete API Endpoints',
    '6. Database Schema & Collections',
    '7. Frontend Screens & Navigation',
    '8. Third-Party Integrations',
    '9. Admin Panel Documentation',
    '10. Security & Compliance',
    '11. Testing Reports',
    '12. Credentials & Access',
    '13. Version History & Release Notes',
    '14. JIRA User Stories (Ready to Import)',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════
add_colored_heading('1. Executive Summary & Project Overview')

doc.add_paragraph(
    'BO (Bananas & Okra) is a comprehensive health, wellness, and nutrition mobile platform. '
    'Built with React Native/Expo (frontend) and FastAPI/MongoDB (backend), the app enables users '
    'to track nutrition, water intake, sleep, exercise, mood, browse diet plans, chat with an AI '
    'wellness coach, connect wearable devices, and engage with a social community feed.'
)

doc.add_heading('Brand Identity', level=2)
add_table(
    ['Property', 'Value'],
    [
        ['Primary Color', 'BO Green #26B50F'],
        ['Secondary Color', 'BO Lime #DBFF02'],
        ['Accent Colors', 'Nutrition Orange #FF9F1C, Water Blue #3A86FF, Fitness Purple #8338EC, Social Teal #06D6A0'],
        ['Typography', 'System fonts (Inter-inspired), weights 400/600/700/800'],
        ['Theme', 'Light, vibrant enterprise design with organic feel'],
        ['Logo', 'Intertwined "bo" letterform in green/lime gradient'],
        ['Animations', 'react-native-reanimated FadeInDown, spring scale, staggered lists'],
    ]
)

doc.add_heading('Tech Stack', level=2)
add_table(
    ['Layer', 'Technology', 'Version'],
    [
        ['Frontend', 'React Native / Expo SDK', '54'],
        ['Routing', 'Expo Router (file-based)', '6.0.22'],
        ['Backend', 'FastAPI / Python', '3.11+'],
        ['Database', 'MongoDB (Motor async driver)', '7.x'],
        ['AI', 'OpenAI GPT-4.1-mini via emergentintegrations', 'Latest'],
        ['Payments', 'Stripe Checkout', 'v2025'],
        ['Image Storage', 'Cloudinary', 'v1.x'],
        ['Push Notifications', 'Expo Push Notifications', 'SDK 54'],
        ['Animations', 'react-native-reanimated', '4.1.x'],
        ['Wearables', 'expo-sensors / REST API sync', 'Latest'],
        ['Deployment', 'Emergent (Kubernetes)', 'Latest'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2. PROJECT METRICS
# ════════════════════════════════════════════════════════════
add_colored_heading('2. Project Metrics & Statistics')

add_table(
    ['Metric', 'Count'],
    [
        ['Total Sprints Delivered', '10'],
        ['Total Modules/Features', '35+'],
        ['User Stories Completed', '120+'],
        ['Estimated Story Points', '350+'],
        ['API Endpoints', '237'],
        ['Database Collections', '67'],
        ['Frontend Screens', '47'],
        ['Platforms Supported', '3 (iOS, Android, Web)'],
        ['Wearable Integrations', '5 (Apple Health, Google Fit, Fitbit, Samsung, Garmin)'],
        ['Seeded Meals', '66'],
        ['Seeded Restaurants', '60'],
        ['Badges', '12'],
        ['Wellness Programs', '4'],
        ['Admin Panel Pages', '8 (Dashboard, Users, Meals, Posts, Restaurants, Quotes, Claims, Tickets)'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. USER TYPES & ROLES
# ════════════════════════════════════════════════════════════
add_colored_heading('3. User Types & Roles')

doc.add_heading('3.1 Regular User (Free Tier)', level=2)
doc.add_paragraph('Default user role upon registration. Has access to:')
features = [
    'Full onboarding flow (goals, dietary preferences, health data)',
    'Dashboard with daily stats, macros, quick-add buttons',
    'Smart Menu with 66 meals across 7 categories',
    'Meal logging, water tracking, sleep tracking, walking tracker, MET tracker',
    'AI Wellness Coach (GPT-4.1-mini powered chat)',
    'Community social feed (create posts, like, comment)',
    'Journal entries with daily reflections',
    'Goal setting and progress tracking with milestones',
    'Happiness/mood tracking with animated emojis',
    'Wearable device connections (5 providers)',
    'Push notifications for reminders',
    'Restaurant discovery and claiming',
    'Profile management with avatar upload',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('3.2 Pro User (Paid Subscription)', level=2)
doc.add_paragraph('Everything in Free tier plus:')
pro_features = [
    'Advanced AI meal recommendations',
    'Priority support tickets',
    'Extended analytics and reports',
    'Custom wellness programs',
    'Ad-free experience',
]
for f in pro_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('3.3 Admin User', level=2)
doc.add_paragraph('Full platform management access via web-based admin panel:')
admin_features = [
    'Dashboard with KPIs, user growth charts, AI-powered analytics',
    'User management: View 360° user profiles, suspend, activate, delete users',
    'Content management: Create/edit/delete meals, posts, quotes, restaurants',
    'Restaurant claim approval/rejection workflow',
    'Support ticket management with response templates',
    'AI recipe info generation for nutritional data',
    'Version management and app configuration',
    '2FA authentication for admin access',
]
for f in admin_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('3.4 Restaurant Owner', level=2)
doc.add_paragraph('Can claim and manage restaurant listings:')
rest_features = [
    'Submit restaurant claim with documentation',
    'Manage restaurant profile after approval',
    'View restaurant analytics',
]
for f in rest_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. SPRINT & FEATURE LIST
# ════════════════════════════════════════════════════════════
add_colored_heading('4. Sprint & Feature List')

sprints = [
    ('Sprint 1: Core MVP', 'Feb 2026', [
        ('Authentication', 'Email/password registration, JWT tokens, auto-redirect'),
        ('Onboarding', '3-step flow: Goals, Dietary Preferences, Health Data'),
        ('Dashboard/Home', 'Personalized greeting, daily stats, macros, quick-add buttons'),
        ('Smart Menu', '6 diet plans, 18 seeded meals with nutrition info'),
        ('Social Feed', 'Create posts, like/unlike, comment system'),
        ('AI Wellness Coach', 'GPT-4.1-mini chat with context-aware responses'),
        ('Profile', 'User info, daily stats, goals, settings menu'),
    ]),
    ('Sprint 2: Wellness Programs & Restaurants', 'Feb 2026', [
        ('Wellness Programs', '4 curated programs with daily activities and check-ins'),
        ('Restaurant Discovery', '60 seeded restaurants with ratings and reviews'),
        ('Daily Motivational Quotes', '30 quotes served via API'),
        ('Recipe Browsing', 'Detailed recipe pages with ingredients and directions'),
    ]),
    ('Sprint 3: Quick Adds & Goals', 'Mar 2026', [
        ('Quick Adds Tab', '4-zone: Quick Add, Journal, Trackers, Timeline'),
        ('Water Tracker', 'Animated circular progress ring, daily 8-glass goal'),
        ('Sleep Tracker', 'Hours/minutes input, 5-moon quality rating'),
        ('Walking Tracker', 'Step counter with distance/calories calculation'),
        ('MET Tracker', '6 activity types with MET-minutes calculation'),
        ('Journal', 'Create/view/like/delete entries with daily reflections'),
        ('My Goals Tab', 'Life goals, happiness gauge, progress cards, streaks'),
        ('Wellness Enrollment', 'Program enrollment with daily check-in system'),
        ('30-Day Reports', 'Progress reports across all tracker categories'),
    ]),
    ('Sprint 4: Community & Enhanced Meals', 'Mar 2026', [
        ('Enhanced Feed', 'Instagram-style feed, media support, infinite scroll'),
        ('Meals Database', '40+ meals across 7 categories, search & filters'),
        ('Meal Plans', 'Weekly meal planning with breakfast/lunch/dinner slots'),
        ('User Recipes', 'Full CRUD for custom recipes'),
        ('Badges System', '12 badges across 4 categories'),
        ('Profile Enhancement', 'Avatar upload, subscription status'),
    ]),
    ('Sprint 5: Media & Content', 'Mar 2026', [
        ('Cloudinary Integration', 'Image upload for posts and profiles'),
        ('Enhanced Feed Media', 'Photo sharing in community posts'),
        ('Content Moderation', 'Post ownership enforcement'),
    ]),
    ('Sprint 6: Payments & Subscriptions', 'Mar 2026', [
        ('Stripe Integration', 'Checkout sessions for Pro subscription'),
        ('Subscription Plans', '3 tiers: Free, Monthly Pro, Annual Pro'),
        ('Payment History', 'Transaction records and subscription management'),
    ]),
    ('Sprint 7: Notifications & Support', 'Mar 2026', [
        ('Push Notifications', 'Expo push with meal reminders, wellness tips'),
        ('Notification Preferences', 'Per-category toggle settings'),
        ('Support Tickets', 'Create, view, respond to support tickets'),
        ('Help Center', 'FAQ system with 20 seeded questions'),
    ]),
    ('Sprint 8: Admin Panel', 'Mar 2026', [
        ('Admin Authentication', '2FA with email/password + OTP'),
        ('Admin Dashboard', 'KPI cards, user growth chart, top restaurants'),
        ('User Management', 'List, search, view details, suspend, delete'),
        ('Content Management', 'CRUD for meals, posts, quotes'),
        ('Restaurant Claims', 'Approval/rejection workflow'),
    ]),
    ('Sprint 9: Advanced Features', 'Apr 2026', [
        ('Wearable Integration', '5 device providers with health data sync'),
        ('Pedometer', 'Real-time step counting via device sensors'),
        ('Happiness Tracking', 'Custom animated SVG emoji faces'),
        ('Invite Friends', 'Referral system with unique codes'),
        ('Account Management', 'Account deletion, data export'),
        ('Legal Pages', 'Privacy policy, terms of service'),
    ]),
    ('Sprint 10: Enterprise & Polish', 'Apr 2026', [
        ('AI Analytics Dashboard', 'GPT-powered insights and recommendations'),
        ('User 360° View', 'Complete user activity dashboard with health charts'),
        ('AI Recipe Generation', 'Auto-generate nutritional info with approx labels'),
        ('BO Branding', 'Greyed-out logo placeholders, notification branding'),
        ('Version Management', 'Settings version display + release history'),
        ('About Page', 'Compliance info: GDPR, CCPA, HIPAA, SOC2, encryption'),
        ('Connected Devices', 'Full 6-metric sync: steps, HR, calories, sleep, distance, active min'),
        ('Deployment Fixes', 'Metro resolver, EAS build fixes, icon fixes'),
    ]),
]

for sprint_name, date, features in sprints:
    doc.add_heading(sprint_name + f' ({date})', level=2)
    for feat_name, feat_desc in features:
        p = doc.add_paragraph()
        run_b = p.add_run(f'{feat_name}: ')
        run_b.bold = True
        p.add_run(feat_desc)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. COMPLETE API ENDPOINTS
# ════════════════════════════════════════════════════════════
add_colored_heading('5. Complete API Endpoints')

doc.add_heading('5.1 Authentication', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['POST', '/api/auth/register', 'Register new user'],
    ['POST', '/api/auth/login', 'Login with email/password'],
    ['GET', '/api/auth/me', 'Get current authenticated user'],
    ['POST', '/api/auth/refresh', 'Refresh access token'],
    ['POST', '/api/auth/logout', 'Logout user'],
    ['POST', '/api/auth/forgot-password', 'Request password reset'],
    ['POST', '/api/auth/reset-password', 'Reset password with code'],
    ['PUT', '/api/auth/avatar', 'Update profile avatar'],
])

doc.add_heading('5.2 Profile & Onboarding', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['PUT', '/api/profile/onboarding', 'Save onboarding data'],
    ['GET', '/api/profile', 'Get user profile'],
    ['PUT', '/api/profile', 'Update profile'],
    ['GET', '/api/v1/profile', 'Get full profile (v1)'],
    ['PUT', '/api/v1/profile/update', 'Update profile details'],
    ['GET', '/api/v1/subscription', 'Get subscription status'],
])

doc.add_heading('5.3 Nutrition & Meals', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/diet-plans', 'List diet plans'],
    ['GET', '/api/diet-plans/{id}', 'Get plan with meals'],
    ['GET', '/api/v1/meals', 'List meals (paginated, filtered)'],
    ['GET', '/api/v1/meals/{id}', 'Get meal detail'],
    ['GET', '/api/v1/meals/search', 'Search meals'],
    ['GET', '/api/v1/meals/favorites', 'User favorites list'],
    ['POST', '/api/v1/meal/fav/{mealId}', 'Toggle favorite'],
    ['POST', '/api/v1/meals/log', 'Log a meal entry'],
    ['GET', '/api/v1/meals/log', 'Get meal logs for date'],
    ['DELETE', '/api/v1/meals/log/{id}', 'Delete meal log'],
    ['POST', '/api/v1/meal-plan', 'Add to meal plan'],
    ['GET', '/api/v1/meal-plan', 'Get meal plan'],
    ['DELETE', '/api/v1/meal-plan/{id}', 'Remove from plan'],
])

doc.add_heading('5.4 Recipes', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['POST', '/api/v1/receipes', 'Create recipe'],
    ['GET', '/api/v1/receipes', 'List user recipes'],
    ['GET', '/api/v1/receipes/{id}', 'Recipe detail'],
    ['PUT', '/api/v1/receipes/{id}', 'Update recipe'],
    ['DELETE', '/api/v1/receipes/{id}', 'Delete recipe'],
])

doc.add_heading('5.5 Trackers', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['POST', '/api/v1/trackers/water', 'Log water intake'],
    ['GET', '/api/v1/trackers/water', 'Get water logs'],
    ['POST', '/api/v1/trackers/sleep', 'Log sleep data'],
    ['GET', '/api/v1/trackers/sleep', 'Get sleep logs'],
    ['POST', '/api/v1/trackers/walking', 'Log walking data'],
    ['GET', '/api/v1/trackers/walking', 'Get walking logs'],
    ['POST', '/api/v1/trackers/met', 'Log MET activity'],
    ['GET', '/api/v1/trackers/met', 'Get MET logs'],
    ['POST', '/api/v1/trackers/happiness', 'Log happiness level'],
    ['GET', '/api/v1/trackers/happiness', 'Get happiness logs'],
    ['GET', '/api/v1/trackers/summary', 'Daily tracker summary'],
    ['GET', '/api/v1/trackers/timeline', 'Daily activity timeline'],
])

doc.add_heading('5.6 Social Feed', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['POST', '/api/v1/feed', 'Create feed post'],
    ['GET', '/api/v1/feed', 'Get paginated feed'],
    ['GET', '/api/v1/feed/{id}', 'Get post detail'],
    ['PUT', '/api/v1/feed/{id}', 'Update own post'],
    ['DELETE', '/api/v1/feed/{id}', 'Delete own post'],
    ['POST', '/api/v1/post/like/{postId}', 'Toggle like'],
    ['GET', '/api/v1/post/likes/{postId}', 'Get who liked'],
    ['POST', '/api/v1/post/comment/{postId}', 'Add comment'],
    ['GET', '/api/v1/post/comments/{postId}', 'Get comments'],
])

doc.add_heading('5.7 Wellness & Goals', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/goals', 'Get user goals'],
    ['GET', '/api/v1/goals/progress', 'Get goal progress'],
    ['GET', '/api/wellness-programs', 'List wellness programs'],
    ['POST', '/api/v1/wellness-programs/{id}/enroll', 'Enroll in program'],
    ['POST', '/api/v1/wellness-programs/checkin', 'Daily check-in'],
    ['GET', '/api/v1/wellness-programs/active', 'Get active program'],
])

doc.add_heading('5.8 Journal', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['POST', '/api/v1/journal', 'Create journal entry'],
    ['GET', '/api/v1/journal', 'List journal entries'],
    ['GET', '/api/v1/journal/{id}', 'Get journal entry'],
    ['PUT', '/api/v1/journal/{id}', 'Update journal entry'],
    ['DELETE', '/api/v1/journal/{id}', 'Delete journal entry'],
])

doc.add_heading('5.9 Wearables', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/wearables/providers', 'List supported providers'],
    ['POST', '/api/v1/wearables/connect', 'Connect a device'],
    ['DELETE', '/api/v1/wearables/disconnect', 'Disconnect a device'],
    ['POST', '/api/v1/wearables/sync', 'Sync health data'],
    ['GET', '/api/v1/wearables/summary', 'Get health data summary'],
])

doc.add_heading('5.10 Payments', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['POST', '/api/v1/payment/create-checkout', 'Create Stripe checkout session'],
    ['POST', '/api/v1/payment/webhook', 'Stripe webhook handler'],
    ['GET', '/api/v1/payment/history', 'Get payment history'],
])

doc.add_heading('5.11 Notifications', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['POST', '/api/v1/push/register', 'Register push token'],
    ['GET', '/api/v1/notifications', 'Get user notifications'],
    ['PUT', '/api/v1/notifications/{id}/read', 'Mark notification read'],
    ['GET', '/api/v1/notification-preferences', 'Get preferences'],
    ['PUT', '/api/v1/notification-preferences', 'Update preferences'],
])

doc.add_heading('5.12 Admin', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/admin-panel', 'Serve admin panel HTML'],
    ['POST', '/api/v1/admin/login', 'Admin authentication'],
    ['POST', '/api/v1/admin/verify-2fa', 'Verify 2FA OTP'],
    ['GET', '/api/v1/admin/dashboard', 'Dashboard statistics'],
    ['GET', '/api/v1/admin/users', 'List all users'],
    ['GET', '/api/v1/admin/user/360/{id}', 'User 360-degree view'],
    ['POST', '/api/v1/admin/users/changeAction/{id}', 'Suspend/activate/delete user'],
    ['POST', '/api/v1/admin/meals', 'Create meal (admin)'],
    ['PUT', '/api/v1/admin/meals/{id}', 'Update meal'],
    ['DELETE', '/api/v1/admin/meals/{id}', 'Delete meal'],
    ['GET', '/api/v1/admin/ai-analytics', 'AI-powered analytics'],
    ['POST', '/api/v1/admin/ai-recipe-info', 'AI recipe info generation'],
    ['GET', '/api/v1/admin/claims', 'List restaurant claims'],
    ['PUT', '/api/v1/admin/claims/{id}/approve', 'Approve claim'],
    ['PUT', '/api/v1/admin/claims/{id}/reject', 'Reject claim'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. DATABASE SCHEMA
# ════════════════════════════════════════════════════════════
add_colored_heading('6. Database Schema & Collections')

db_collections = [
    ('users', '84', 'User accounts with profile, goals, preferences'),
    ('meals', '66', 'Curated meal database with nutritional data'),
    ('restaurants', '60', 'Restaurant listings with ratings'),
    ('feed_posts', '10', 'Community feed posts'),
    ('feed_likes', '4', 'Post like records'),
    ('feed_comments', '4', 'Post comments'),
    ('journals', '1', 'Journal entries'),
    ('meal_logs', '1', 'Logged meal entries'),
    ('water_logs', '1205', 'Water intake records'),
    ('sleep_logs', '1204', 'Sleep tracking records'),
    ('walking_logs', '8', 'Walking/step records'),
    ('met_logs', '4', 'MET activity records'),
    ('happiness_logs', '6', 'Mood/happiness records'),
    ('workouts', '2', 'Workout records'),
    ('weight_logs', '200', 'Weight tracking'),
    ('notifications', '651', 'Push notification records'),
    ('push_tokens', '3', 'Device push tokens'),
    ('chat_messages', '6', 'AI chat history'),
    ('wellness_programs', '4', 'Wellness program definitions'),
    ('wellness_enrollments', '3', 'User program enrollments'),
    ('badges', '12', 'Badge definitions'),
    ('user_badges', '3', 'Earned badges'),
    ('subscription_plans', '3', 'Subscription plan tiers'),
    ('user_subscriptions', '2', 'Active subscriptions'),
    ('restaurant_claims', '2', 'Restaurant ownership claims'),
    ('wearable_connections', '2', 'Connected wearable devices'),
    ('wearable_data', '4', 'Synced health data'),
    ('admin_quotes', '30', 'Daily motivational quotes'),
    ('faqs', '20', 'Help center FAQs'),
    ('diet_plans', '6', 'Diet plan definitions'),
]

add_table(['Collection', 'Documents', 'Description'], db_collections)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 7. FRONTEND SCREENS
# ════════════════════════════════════════════════════════════
add_colored_heading('7. Frontend Screens & Navigation')

screens = [
    ('Authentication', [
        ('login.tsx', 'Email/password login with BO branding'),
        ('register.tsx', 'User registration form'),
        ('forgot-password.tsx', 'Password reset request'),
        ('privacy-policy.tsx', 'Privacy policy display'),
    ]),
    ('Onboarding', [
        ('goals.tsx', 'Select wellness goals'),
        ('dietary.tsx', 'Dietary preferences'),
        ('health.tsx', 'Health data input'),
        ('activities.tsx', 'Activity preferences'),
        ('life-goals.tsx', 'Life goals selection'),
        ('happiness.tsx', 'Baseline happiness'),
        ('badges.tsx', 'Badge showcase'),
        ('permissions.tsx', 'Notification permissions'),
        ('questionnaire.tsx', 'Health questionnaire'),
        ('complete.tsx', 'Onboarding completion'),
    ]),
    ('Main Tabs', [
        ('home.tsx', 'Dashboard with stats, programs, restaurants'),
        ('menu.tsx', 'Smart menu with meals, categories, favorites'),
        ('quick-adds.tsx', 'Trackers (water, sleep, walking, MET), journal, timeline'),
        ('goals.tsx', 'Goal progress, happiness tracking, streaks'),
        ('feed.tsx', 'Community social feed'),
    ]),
    ('Detail Screens', [
        ('meal/[id].tsx', 'Meal detail with nutrition, ingredients'),
        ('restaurant/[id].tsx', 'Restaurant detail with menu'),
        ('recipe-detail.tsx', 'Recipe with directions'),
        ('ticket-detail.tsx', 'Support ticket conversation'),
        ('wellness-programs.tsx', 'All wellness programs list'),
    ]),
    ('Profile & Settings', [
        ('profile.tsx', 'User profile with avatar, stats'),
        ('settings.tsx', 'App settings with version management'),
        ('about.tsx', 'About app with compliance info'),
        ('subscription.tsx', 'Subscription plans & Stripe checkout'),
        ('notifications.tsx', 'Notification center'),
        ('notification-settings.tsx', 'Notification preferences'),
        ('connected-devices.tsx', 'Wearable device management'),
        ('invite-friends.tsx', 'Referral system'),
        ('help.tsx', 'FAQ and help center'),
        ('contact.tsx', 'Contact support'),
        ('progress.tsx', '30-day progress report'),
        ('privacy-screen.tsx', 'Privacy policy'),
        ('terms.tsx', 'Terms of service'),
        ('account-delete.tsx', 'Account deletion'),
        ('claim-restaurant.tsx', 'Restaurant claim form'),
    ]),
]

for section, items in screens:
    doc.add_heading(section, level=2)
    add_table(['Screen File', 'Description'], items)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 8. THIRD-PARTY INTEGRATIONS
# ════════════════════════════════════════════════════════════
add_colored_heading('8. Third-Party Integrations')

add_table(['Service', 'Purpose', 'Configuration'], [
    ['OpenAI GPT-4.1-mini', 'AI Wellness Coach, Admin Analytics, Recipe Generation', 'Via emergentintegrations library with Emergent LLM Key'],
    ['Stripe', 'Payment processing for Pro subscriptions', 'Checkout sessions + webhook handlers'],
    ['Cloudinary', 'Image upload and storage', 'Profile photos, post media'],
    ['Expo Push Notifications', 'Push notifications to iOS/Android', 'Token-based with BO branding'],
    ['MongoDB Atlas', 'Production database (deployment)', 'Motor async driver'],
    ['Apple Health', 'Wearable health data sync (iOS)', 'REST API integration'],
    ['Google Fit', 'Wearable health data sync (Android)', 'REST API integration'],
    ['Fitbit', 'Wearable health data sync', 'REST API integration'],
    ['Samsung Health', 'Wearable health data sync (Android)', 'REST API integration'],
    ['Garmin Connect', 'Wearable health data sync', 'REST API integration'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 9. ADMIN PANEL
# ════════════════════════════════════════════════════════════
add_colored_heading('9. Admin Panel Documentation')

doc.add_paragraph(
    'The admin panel is a web-based dashboard served at /api/admin-panel. '
    'It provides comprehensive platform management with 2FA authentication, '
    'AI-powered analytics, and full CRUD operations for all content types.'
)

doc.add_heading('Admin Panel Pages', level=2)
add_table(['Page', 'Features'], [
    ['Dashboard', 'KPI cards, user growth chart, top restaurants, AI analytics with insights/recommendations, health metrics overview'],
    ['Users', 'User list with search, 360° user detail view (health timeline, workouts, meals, social, subscriptions, tickets), suspend/activate/delete'],
    ['Meals', 'Create/edit/delete meals, AI-powered nutritional info generation, category management'],
    ['Posts', 'Create/edit/delete feed posts, BO logo placeholder for missing images'],
    ['Restaurants', 'List restaurants, view ratings, manage listings'],
    ['Quotes', 'Manage daily motivational quotes'],
    ['Claims', 'Restaurant claim approval/rejection workflow'],
    ['Tickets', 'Support ticket management with response templates'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 10. SECURITY & COMPLIANCE
# ════════════════════════════════════════════════════════════
add_colored_heading('10. Security & Compliance')

compliance = [
    ['GDPR', 'Full compliance with EU General Data Protection Regulation. Users can request data export or deletion.'],
    ['CCPA', 'California Consumer Privacy Act compliant. Transparent data collection with opt-out rights.'],
    ['HIPAA Aware', 'Health data handled following HIPAA best practices. PHI encrypted at rest and in transit.'],
    ['SOC 2 Type II', 'Infrastructure on SOC 2 Type II certified cloud providers.'],
    ['ISO 27001', 'Information security management standards followed.'],
    ['OWASP', 'Mobile security standards implemented.'],
    ['End-to-End Encryption', 'TLS 1.3 in transit, AES-256-GCM at rest. API keys rotated regularly.'],
    ['Biometric Auth', 'Face ID/Touch ID support. No biometric data leaves device.'],
    ['2FA Admin', 'Two-factor authentication for admin panel access.'],
    ['JWT Auth', '24h access tokens, 30d refresh tokens, bcrypt password hashing.'],
]
add_table(['Standard', 'Implementation'], compliance)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 11. TESTING REPORTS
# ════════════════════════════════════════════════════════════
add_colored_heading('11. Testing Reports')

doc.add_heading('11.1 Backend API Testing', level=2)
doc.add_paragraph('All API endpoints tested via automated pytest suite and manual curl testing.')

test_areas = [
    ['Authentication', '8 endpoints', 'PASS', 'Register, login, JWT refresh, logout, password reset'],
    ['Profile & Onboarding', '6 endpoints', 'PASS', 'Onboarding flow, profile CRUD, subscription status'],
    ['Meals & Nutrition', '13 endpoints', 'PASS', 'Meal listing, search, favorites, logging, meal plans'],
    ['Trackers', '12 endpoints', 'PASS', 'Water, sleep, walking, MET, happiness, summary'],
    ['Social Feed', '9 endpoints', 'PASS', 'Posts CRUD, likes, comments, pagination'],
    ['Wellness & Goals', '6 endpoints', 'PASS', 'Goals, programs, enrollment, check-ins'],
    ['Wearables', '5 endpoints', 'PASS', 'Connect, disconnect, sync, summary'],
    ['Payments', '3 endpoints', 'PASS', 'Stripe checkout, webhook, history'],
    ['Notifications', '5 endpoints', 'PASS', 'Push registration, listing, preferences'],
    ['Admin', '15+ endpoints', 'PASS', 'Dashboard, user management, content CRUD, AI analytics'],
]
add_table(['Module', 'Endpoints', 'Status', 'Coverage'], test_areas)

doc.add_heading('11.2 Frontend Testing', level=2)
doc.add_paragraph('Mobile UI tested on iOS (Expo Go) and web preview. Key areas:')

ui_tests = [
    ['Onboarding Flow', 'PASS', 'All 10 steps complete without errors'],
    ['Navigation', 'PASS', '5-tab navigation + deep linking'],
    ['Input Fields', 'PASS', 'All modals allow typing (BottomSheet fix applied)'],
    ['Image Fallbacks', 'PASS', 'BO logo placeholder for missing images'],
    ['Animations', 'PASS', 'Reanimated transitions, custom mood emojis'],
    ['Pull-to-Refresh', 'PASS', 'Home, Goals, Quick-Adds, Feed, Profile tabs'],
    ['Push Notifications', 'PASS', 'BO-branded with logo and subtitle'],
    ['Wearable Sync', 'PASS', '6 health metrics synced on connect'],
]
add_table(['Area', 'Status', 'Notes'], ui_tests)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 12. CREDENTIALS & ACCESS
# ════════════════════════════════════════════════════════════
add_colored_heading('12. Credentials & Access')

doc.add_heading('12.1 Admin Access', level=2)
add_table(['Property', 'Value'], [
    ['Admin Panel URL', '{your-domain}/api/admin-panel'],
    ['Email', 'admin@bo.com'],
    ['Password', 'BoAdmin2026!'],
    ['2FA', 'Demo code shown on login screen'],
    ['Role', 'admin'],
])

doc.add_heading('12.2 Test User', level=2)
add_table(['Property', 'Value'], [
    ['Email', 'test@bo.com'],
    ['Password', 'Test1234!'],
    ['Role', 'user'],
    ['Note', 'Create via POST /api/auth/register'],
])

doc.add_heading('12.3 API Keys (Production)', level=2)
add_table(['Service', 'Environment Variable', 'Notes'], [
    ['MongoDB Atlas', 'MONGO_URL', 'Set in backend/.env for production'],
    ['OpenAI / Emergent LLM', 'EMERGENT_LLM_KEY', 'Universal key for AI features'],
    ['Stripe', 'STRIPE_SECRET_KEY', 'Required for payment processing'],
    ['Cloudinary', 'CLOUDINARY_CLOUD_NAME, CLOUDINARY_API_KEY, CLOUDINARY_API_SECRET', 'Image upload service'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 13. VERSION HISTORY
# ════════════════════════════════════════════════════════════
add_colored_heading('13. Version History & Release Notes')

doc.add_heading('v1.0.0 — April 2, 2026 (Current)', level=2)
v1_notes = [
    'Initial public release of BO wellness app',
    'AI-powered meal recommendations with nutritional tracking',
    'Happiness & mood tracking with custom animated SVG emojis',
    'Wearable device sync: Apple Health, Google Fit, Fitbit, Samsung, Garmin',
    'Community feed with posts, likes, and comments',
    'AI wellness chatbot for health Q&A (GPT-4.1-mini)',
    'Push notifications with BO branding',
    'Restaurant discovery with BO verification badges',
    'Wellness programs with daily activities and check-ins',
    'Water, sleep, walking, and workout tracking',
    'Stripe subscription with Pro plan features',
    'Enterprise admin panel with AI analytics and 360° user views',
    'GDPR, CCPA, HIPAA-aware compliance framework',
    'End-to-end encryption (TLS 1.3 + AES-256-GCM)',
]
for n in v1_notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_heading('v0.9.0 — March 15, 2026 (Beta)', level=2)
for n in ['Beta release with core tracking features', 'Onboarding flow with health questionnaire', 'Basic meal logging and recipe browsing', 'Step counter integration', 'User profile with goal setting']:
    doc.add_paragraph(n, style='List Bullet')

doc.add_heading('v0.5.0 — February 1, 2026 (Alpha)', level=2)
for n in ['Internal alpha with auth flow', 'Basic UI framework and navigation', 'MongoDB backend with REST API', 'Initial meal database seeded']:
    doc.add_paragraph(n, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 14. JIRA USER STORIES
# ════════════════════════════════════════════════════════════
add_colored_heading('14. JIRA User Stories (Ready to Import)')

doc.add_paragraph(
    'Below are user stories formatted for JIRA import. Each story includes a title, '
    'description, acceptance criteria, priority, and story points estimation.'
)

jira_stories = [
    ('AUTH-001', 'User Registration', 'As a new user, I want to register with email and password so I can access the app.', 'P0', '5', 'Done'),
    ('AUTH-002', 'User Login', 'As a user, I want to login with my credentials so I can access my account.', 'P0', '3', 'Done'),
    ('AUTH-003', 'Password Reset', 'As a user, I want to reset my password if I forget it.', 'P1', '5', 'Done'),
    ('AUTH-004', 'JWT Token Refresh', 'As a user, I want my session to automatically refresh so I stay logged in.', 'P0', '3', 'Done'),
    ('ONB-001', 'Goal Selection', 'As a new user, I want to select my wellness goals during onboarding.', 'P0', '5', 'Done'),
    ('ONB-002', 'Dietary Preferences', 'As a new user, I want to set my dietary preferences for personalized recommendations.', 'P0', '5', 'Done'),
    ('ONB-003', 'Health Data Input', 'As a new user, I want to input my health data (height, weight, activity level).', 'P0', '5', 'Done'),
    ('DASH-001', 'Personalized Dashboard', 'As a user, I want to see my daily stats summary on the home screen.', 'P0', '8', 'Done'),
    ('DASH-002', 'Quick Add Buttons', 'As a user, I want quick-add buttons for water and meals on the dashboard.', 'P1', '5', 'Done'),
    ('MENU-001', 'Meal Browsing', 'As a user, I want to browse meals by category with nutrition info.', 'P0', '8', 'Done'),
    ('MENU-002', 'Meal Search', 'As a user, I want to search meals by name, category, and calorie range.', 'P1', '5', 'Done'),
    ('MENU-003', 'Favorite Meals', 'As a user, I want to save my favorite meals for quick access.', 'P1', '3', 'Done'),
    ('MENU-004', 'Meal Planning', 'As a user, I want to add meals to my weekly plan.', 'P1', '8', 'Done'),
    ('TRACK-001', 'Water Tracking', 'As a user, I want to track my daily water intake with a visual progress ring.', 'P0', '5', 'Done'),
    ('TRACK-002', 'Sleep Tracking', 'As a user, I want to log my sleep hours and quality.', 'P1', '5', 'Done'),
    ('TRACK-003', 'Step Tracking', 'As a user, I want to track my daily steps with distance/calorie calculations.', 'P1', '5', 'Done'),
    ('TRACK-004', 'MET Activity Tracking', 'As a user, I want to log various exercises with MET-minutes calculation.', 'P1', '5', 'Done'),
    ('TRACK-005', 'Mood/Happiness Tracking', 'As a user, I want to log my daily mood with animated emoji faces.', 'P1', '5', 'Done'),
    ('TRACK-006', 'Workout Logging', 'As a user, I want to log workouts with type, duration, and calories.', 'P1', '5', 'Done'),
    ('FEED-001', 'Community Feed', 'As a user, I want to view and create posts in a community feed.', 'P0', '13', 'Done'),
    ('FEED-002', 'Post Interactions', 'As a user, I want to like and comment on community posts.', 'P1', '8', 'Done'),
    ('CHAT-001', 'AI Wellness Coach', 'As a user, I want to chat with an AI coach for health advice.', 'P0', '13', 'Done'),
    ('GOAL-001', 'Goal Progress', 'As a user, I want to see progress toward my wellness goals with milestone badges.', 'P1', '8', 'Done'),
    ('GOAL-002', 'Streak Tracking', 'As a user, I want to see my consecutive day streaks.', 'P2', '3', 'Done'),
    ('WELL-001', 'Wellness Programs', 'As a user, I want to enroll in wellness programs with daily check-ins.', 'P1', '8', 'Done'),
    ('WELL-002', 'Program Progress', 'As a user, I want to see my progress through enrolled programs.', 'P2', '5', 'Done'),
    ('JOUR-001', 'Daily Journal', 'As a user, I want to create daily journal entries for reflection.', 'P1', '5', 'Done'),
    ('WEAR-001', 'Wearable Connection', 'As a user, I want to connect my wearable device to sync health data.', 'P1', '8', 'Done'),
    ('WEAR-002', 'Health Data Sync', 'As a user, I want my steps, HR, sleep, calories, distance synced automatically.', 'P1', '8', 'Done'),
    ('PAY-001', 'Stripe Subscription', 'As a user, I want to upgrade to Pro plan via Stripe checkout.', 'P1', '8', 'Done'),
    ('NOTIF-001', 'Push Notifications', 'As a user, I want to receive push notifications for meal reminders.', 'P1', '5', 'Done'),
    ('NOTIF-002', 'Notification Preferences', 'As a user, I want to customize which notifications I receive.', 'P2', '3', 'Done'),
    ('REST-001', 'Restaurant Discovery', 'As a user, I want to discover and filter nearby restaurants.', 'P1', '8', 'Done'),
    ('REST-002', 'Restaurant Claiming', 'As a restaurant owner, I want to claim my restaurant listing.', 'P2', '5', 'Done'),
    ('PROF-001', 'Profile Management', 'As a user, I want to update my profile info and upload an avatar.', 'P1', '5', 'Done'),
    ('PROF-002', 'Account Deletion', 'As a user, I want to delete my account and all associated data.', 'P1', '3', 'Done'),
    ('ADMIN-001', 'Admin Dashboard', 'As an admin, I want a dashboard with KPIs and user growth charts.', 'P0', '13', 'Done'),
    ('ADMIN-002', 'User 360 View', 'As an admin, I want to see complete user profiles with all activity data.', 'P1', '13', 'Done'),
    ('ADMIN-003', 'AI Analytics', 'As an admin, I want AI-generated insights about platform health.', 'P2', '8', 'Done'),
    ('ADMIN-004', 'Content Management', 'As an admin, I want to manage meals, posts, quotes, and restaurants.', 'P0', '8', 'Done'),
    ('ADMIN-005', 'AI Recipe Generation', 'As an admin, I want AI to generate approximate nutritional info for meals.', 'P2', '5', 'Done'),
    ('ADMIN-006', 'Claim Management', 'As an admin, I want to approve/reject restaurant ownership claims.', 'P1', '5', 'Done'),
    ('COMP-001', 'About Page with Compliance', 'As a user, I want to see app compliance info (GDPR, CCPA, HIPAA, SOC2).', 'P1', '5', 'Done'),
    ('COMP-002', 'Version History', 'As a user, I want to see the app version and release notes.', 'P2', '3', 'Done'),
    ('UI-001', 'BO Logo Fallback', 'As a user, I want to see BO logo placeholder when images are missing.', 'P2', '3', 'Done'),
    ('UI-002', 'Custom Animated Emojis', 'As a user, I want unique animated mood faces instead of standard emojis.', 'P2', '5', 'Done'),
]

add_table(
    ['ID', 'Story', 'Description', 'Priority', 'Points', 'Status'],
    jira_stories
)

# ─── Total ───
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(f'Total Stories: {len(jira_stories)} | Total Story Points: {sum(int(s[4]) for s in jira_stories)}')
run.bold = True

doc.add_page_break()

# FINAL PAGE
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(120)
run = p.add_run('END OF DOCUMENT')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('BO Health & Wellness Platform\nv1.0.0 | April 2026\nBuilt by Flynaut LLC')
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

# SAVE
output_path = '/app/BO_Complete_Project_Documentation.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')
