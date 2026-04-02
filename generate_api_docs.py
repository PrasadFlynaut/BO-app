"""Generate comprehensive BO API Documentation in DOCX format — Internal & External APIs"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

# ─── Styles ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

GREEN = RGBColor(0x26, 0xB5, 0x0F)
DARK = RGBColor(0x1A, 0x20, 0x2C)
GREY = RGBColor(0x71, 0x80, 0x96)

def add_colored_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GREEN if level == 1 else DARK
    return h

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="26B50F"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    return t

def add_info(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    r1.font.size = Pt(9)
    r2 = p.add_run(str(value))
    r2.font.size = Pt(9)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F3F4F6"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    return p

# ════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(80)
run = p.add_run('BO')
run.font.size = Pt(72)
run.font.bold = True
run.font.color.rgb = GREEN

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('API Documentation')
run2.font.size = Pt(28)
run2.font.color.rgb = DARK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Internal & External API Reference\nv1.0.0 | June 2025')
run3.font.size = Pt(14)
run3.font.color.rgb = GREY

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.space_before = Pt(30)
run4 = p4.add_run('Built by Flynaut LLC\nConfidential — For Internal Use & JIRA Integration')
run4.font.size = Pt(11)
run4.font.color.rgb = GREY

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════
add_colored_heading('Table of Contents')
toc = [
    'PART A — INTERNAL API REFERENCE',
    '  1. API Architecture & Conventions',
    '  2. Authentication APIs',
    '  3. Onboarding APIs',
    '  4. Profile & User APIs',
    '  5. Dashboard API',
    '  6. Nutrition & Meals APIs',
    '  7. Recipes APIs',
    '  8. Meal Plans APIs',
    '  9. Trackers APIs (Water, Sleep, Walking, MET, Happiness)',
    '  10. Journal APIs',
    '  11. Goals & Wellness Program APIs',
    '  12. Social Feed APIs',
    '  13. AI Wellness Coach (Chat) APIs',
    '  14. Wearable Device APIs',
    '  15. Payment & Subscription APIs',
    '  16. Notifications & Push APIs',
    '  17. Restaurant APIs',
    '  18. Support Tickets & FAQ APIs',
    '  19. Referrals & Account Management APIs',
    '  20. Admin Panel APIs',
    '  21. Admin Enterprise APIs (AI Analytics, 360 View)',
    '  22. Utility & Health Check APIs',
    '',
    'PART B — EXTERNAL API & THIRD-PARTY INTEGRATIONS',
    '  23. OpenAI (GPT-4.1-mini) via Emergent Integrations',
    '  24. Stripe Payment Gateway',
    '  25. Cloudinary Image CDN',
    '  26. Expo Push Notification Service',
    '  27. MongoDB (Motor Async Driver)',
    '  28. Wearable Provider APIs (Apple Health, Google Fit, etc.)',
    '',
    'PART C — APPENDICES',
    '  29. Environment Variables Reference',
    '  30. Error Code Reference',
    '  31. Rate Limiting & Security',
    '  32. Credentials & Access',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)
        if item.startswith('PART'):
            r.font.bold = True
            r.font.color.rgb = GREEN

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# PART A: INTERNAL API REFERENCE
# ════════════════════════════════════════════════════════════
ph = doc.add_paragraph()
ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
ph.space_before = Pt(40)
rh = ph.add_run('PART A')
rh.font.size = Pt(36)
rh.font.bold = True
rh.font.color.rgb = GREEN
ph2 = doc.add_paragraph()
ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rh2 = ph2.add_run('Internal API Reference')
rh2.font.size = Pt(20)
rh2.font.color.rgb = DARK
doc.add_page_break()

# ─── 1. API Architecture ───
add_colored_heading('1. API Architecture & Conventions')

doc.add_heading('1.1 Base URL', level=2)
doc.add_paragraph('All internal APIs are served under the /api prefix. The backend binds at 0.0.0.0:8001 and the Kubernetes ingress proxies /api/* requests to port 8001.')
add_code_block('Production:  https://{your-domain}/api/*\nDevelopment: http://localhost:8001/api/*')

doc.add_heading('1.2 Authentication', level=2)
doc.add_paragraph('BO uses JSON Web Tokens (JWT) for authentication:')
auth_details = [
    ['Access Token', 'Bearer token in Authorization header', '24 hours', 'HS256'],
    ['Refresh Token', 'Sent as response body', '30 days', 'HS256'],
    ['Password Hashing', 'bcrypt with auto-salt', 'N/A', 'bcrypt'],
]
add_table(['Component', 'Mechanism', 'Expiry', 'Algorithm'], auth_details)
doc.add_paragraph()
doc.add_paragraph('All authenticated endpoints require:')
add_code_block('Authorization: Bearer <access_token>')

doc.add_heading('1.3 Request/Response Format', level=2)
doc.add_paragraph('All requests and responses use JSON (application/json). Dates use ISO 8601 format (YYYY-MM-DDTHH:MM:SS.ffffffZ).')

doc.add_heading('1.4 Pagination Pattern', level=2)
doc.add_paragraph('Paginated endpoints return:')
add_code_block('{\n  "data": [...],\n  "pagination": {\n    "page": 1,\n    "limit": 10,\n    "total": 142,\n    "totalPages": 15,\n    "hasNext": true,\n    "hasPrev": false\n  }\n}')

doc.add_heading('1.5 Error Response Format', level=2)
add_code_block('{\n  "detail": "Human-readable error message"\n}')
doc.add_paragraph('Standard HTTP status codes: 200 (OK), 201 (Created), 400 (Bad Request), 401 (Unauthorized), 403 (Forbidden), 404 (Not Found), 422 (Validation Error), 429 (Rate Limited), 500 (Server Error).')

doc.add_page_break()

# ─── 2. Authentication APIs ───
add_colored_heading('2. Authentication APIs')
doc.add_paragraph('Source: server.py — Handles user registration, login, token management, and password operations.')

doc.add_heading('POST /api/auth/register', level=2)
add_info('Description', 'Register a new user account')
add_info('Auth Required', 'No')
add_info('Rate Limit', '5 requests/minute per IP')
doc.add_paragraph('Request Body:')
add_code_block('{\n  "email": "user@example.com",       // Required, unique\n  "password": "SecurePass123!",       // Required, 8+ chars, 1 upper, 1 digit, 1 special\n  "name": "John Doe",                 // Optional\n  "first_name": "John",               // Optional\n  "last_name": "Doe",                 // Optional\n  "phone": "+1234567890",             // Optional\n  "date_of_birth": "1990-01-15"       // Optional\n}')
doc.add_paragraph('Success Response (200):')
add_code_block('{\n  "access_token": "eyJhbGci...",\n  "refresh_token": "eyJhbGci...",\n  "user": {\n    "id": "664f...",\n    "email": "user@example.com",\n    "name": "John Doe",\n    "role": "user",\n    "subscription": "free",\n    "onboarding_complete": false,\n    ...\n  }\n}')
doc.add_paragraph('Error Responses:')
add_table(['Code', 'Detail'], [
    ['400', 'Email already registered'],
    ['400', 'Password must be at least 8 characters with 1 uppercase, 1 digit, 1 special character'],
])

doc.add_heading('POST /api/auth/login', level=2)
add_info('Description', 'Authenticate user with email and password')
add_info('Auth Required', 'No')
doc.add_paragraph('Request Body:')
add_code_block('{\n  "email": "user@example.com",\n  "password": "SecurePass123!"\n}')
doc.add_paragraph('Success Response (200):')
add_code_block('{\n  "access_token": "eyJhbGci...",\n  "refresh_token": "eyJhbGci...",\n  "user": { ... }\n}')
add_table(['Code', 'Detail'], [['401', 'Invalid email or password']])

doc.add_heading('GET /api/auth/me', level=2)
add_info('Description', 'Get current authenticated user profile')
add_info('Auth Required', 'Yes (Bearer Token)')
doc.add_paragraph('Success Response (200):')
add_code_block('{ "user": { "id": "...", "email": "...", "name": "...", ... } }')

doc.add_heading('POST /api/auth/refresh', level=2)
add_info('Description', 'Refresh an expired access token')
add_info('Auth Required', 'No (uses refresh_token)')
doc.add_paragraph('Request Body:')
add_code_block('{ "refresh_token": "eyJhbGci..." }')
doc.add_paragraph('Success Response (200):')
add_code_block('{ "access_token": "...", "refresh_token": "..." }')

doc.add_heading('POST /api/auth/forgot-password', level=2)
add_info('Description', 'Request a password reset code (6-digit, 10-min expiry, max 3 attempts)')
add_info('Auth Required', 'No')
doc.add_paragraph('Request Body:')
add_code_block('{ "email": "user@example.com" }')
doc.add_paragraph('Response: Returns reset code in dev mode. In production, email via SendGrid (not yet integrated).')

doc.add_heading('POST /api/auth/reset-password', level=2)
add_info('Description', 'Reset password with the 6-digit code')
add_info('Auth Required', 'No')
doc.add_paragraph('Request Body:')
add_code_block('{ "email": "user@example.com", "code": "123456", "new_password": "NewPass123!" }')

doc.add_heading('PUT /api/auth/change-password', level=2)
add_info('Description', 'Change password for the logged-in user')
add_info('Auth Required', 'Yes')
doc.add_paragraph('Request Body:')
add_code_block('{ "current_password": "OldPass123!", "new_password": "NewPass456!" }')

doc.add_heading('PUT /api/auth/avatar', level=2)
add_info('Description', 'Update user profile photo / avatar URL')
add_info('Auth Required', 'Yes')
doc.add_paragraph('Request Body:')
add_code_block('{ "avatar_url": "https://res.cloudinary.com/..." }')

doc.add_heading('POST /api/v1/auth/logout', level=2)
add_info('Description', 'Logout and invalidate session (client-side token removal)')
add_info('Auth Required', 'Yes')

doc.add_page_break()

# ─── 3. Onboarding APIs ───
add_colored_heading('3. Onboarding APIs')
doc.add_paragraph('Source: server.py — Multi-step onboarding flow to personalize the user experience.')

onboarding_endpoints = [
    ['PUT', '/api/profile/onboarding', 'Save main onboarding data (goals, diet, health)', 'Yes', '{ goals, dietary_preferences, allergies, height_cm, weight_kg, target_weight_kg, gender, activity_level, sleep_hours }'],
    ['POST', '/api/onboarding/activities', 'Save activity & fitness preferences', 'Yes', '{ activities: [...], fitness_goals: [...] }'],
    ['PUT', '/api/onboarding/preferences', 'Save meal & dietary preferences', 'Yes', '{ meal_preferences: [...], allergies: [...] }'],
    ['PUT', '/api/onboarding/questionnaire', 'Save health questionnaire', 'Yes', '{ favorite_fast_food, dietary_restriction, under_nutritionist, health_info, ... }'],
    ['PUT', '/api/onboarding/life-goals', 'Save life goals & happiness baseline', 'Yes', '{ life_goals: [...], happiness_level: 5, review_text: "" }'],
    ['PUT', '/api/onboarding/permissions', 'Save device permissions', 'Yes', '{ push_notifications, gallery_access, location_sharing, privacy_policy_accepted }'],
    ['POST', '/api/onboarding/complete', 'Mark onboarding as complete', 'Yes', 'No body'],
]
add_table(['Method', 'Endpoint', 'Description', 'Auth', 'Request Body (key fields)'], onboarding_endpoints)

doc.add_page_break()

# ─── 4. Profile APIs ───
add_colored_heading('4. Profile & User APIs')
doc.add_paragraph('Source: server.py, sprint4.py')
profile_eps = [
    ['GET', '/api/profile', 'Get user profile', 'Yes', 'Full user object'],
    ['PUT', '/api/profile', 'Update profile (name, phone, address, bio)', 'Yes', 'Updated user object'],
    ['GET', '/api/v1/profile', 'Get full profile (v1 with extended data)', 'Yes', 'Full user + subscription info'],
    ['PUT', '/api/v1/profile/update', 'Update profile details (v1)', 'Yes', 'Updated user object'],
    ['GET', '/api/v1/subscription', 'Get subscription status', 'Yes', '{ subscription, plan_name, ... }'],
]
add_table(['Method', 'Endpoint', 'Description', 'Auth', 'Response'], profile_eps)

doc.add_page_break()

# ─── 5. Dashboard API ───
add_colored_heading('5. Dashboard API')
doc.add_paragraph('Source: server.py — Aggregates daily stats for the home screen.')
doc.add_heading('GET /api/dashboard', level=2)
add_info('Auth', 'Yes')
doc.add_paragraph('Response:')
add_code_block('{\n  "date": "2026-04-02",\n  "water_ml": 1500, "water_goal_ml": 2500,\n  "calories": 1200, "calorie_goal": 2000,\n  "protein_g": 65, "carbs_g": 120, "fat_g": 40,\n  "meals_logged": 3,\n  "user_name": "John",\n  "weight_kg": 75, "target_weight_kg": 70\n}')

doc.add_page_break()

# ─── 6. Nutrition & Meals APIs ───
add_colored_heading('6. Nutrition & Meals APIs')
doc.add_paragraph('Sources: server.py (basic meals, nutrition logging), sprint4.py (enhanced meals with search, favorites, pagination)')

doc.add_heading('6.1 Diet Plans', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['GET', '/api/diet-plans', 'List all 6 diet plans', 'No'],
    ['GET', '/api/diet-plans/{plan_id}', 'Get plan detail with associated meals', 'No'],
])

doc.add_heading('6.2 Basic Meals', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['GET', '/api/meals', 'List meals (filter: plan_id, category)', 'No'],
    ['GET', '/api/meals/{meal_id}', 'Get single meal detail', 'No'],
])

doc.add_heading('6.3 Enhanced Meals (v1)', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['GET', '/api/v1/meals', 'Paginated meals listing with filters (category, meal_type, sort)', 'Yes'],
    ['GET', '/api/v1/meals/search', 'Search meals by text, category, calorie range, meal type', 'Yes'],
    ['GET', '/api/v1/meals/{meal_id}', 'Get meal with full nutrition, ingredients, directions', 'Yes'],
    ['GET', '/api/v1/meals/favorites', 'List user\'s favorited meals', 'Yes'],
    ['POST', '/api/v1/meal/fav/{meal_id}', 'Toggle meal favorite (like/unlike)', 'Yes'],
])
doc.add_paragraph('Search Query Parameters:')
add_code_block('GET /api/v1/meals/search?q=chicken&category=High+Protein&minCal=200&maxCal=600&mealType=lunch&page=1&limit=10')

doc.add_heading('6.4 Nutrition Logging', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/nutrition/log', 'Log a nutrition entry (meal_type, food_name, calories, macros)', 'Yes'],
    ['GET', '/api/nutrition/daily', 'Get today\'s nutrition logs + totals vs goals', 'Yes'],
])

doc.add_heading('6.5 Meal Logging (v1)', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/v1/meals/log', 'Log a meal entry (meal_type, meal_id, name, calories)', 'Yes'],
    ['GET', '/api/v1/meals/log', 'Get meal logs for a date (?date=2026-04-02)', 'Yes'],
    ['DELETE', '/api/v1/meals/log/{log_id}', 'Delete a meal log entry', 'Yes'],
])

doc.add_page_break()

# ─── 7. Recipes APIs ───
add_colored_heading('7. Recipes APIs')
doc.add_paragraph('Source: sprint4.py — Full CRUD for user-created recipes.')
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/v1/receipes', 'Create a new recipe', 'Yes'],
    ['GET', '/api/v1/receipes', 'List user\'s recipes (paginated)', 'Yes'],
    ['GET', '/api/v1/receipes/{id}', 'Get recipe detail', 'Yes'],
    ['PUT', '/api/v1/receipes/{id}', 'Update recipe (owner only)', 'Yes'],
    ['DELETE', '/api/v1/receipes/{id}', 'Delete recipe (owner only)', 'Yes'],
])
doc.add_paragraph('Create Recipe Request Body:')
add_code_block('{\n  "title": "Grilled Chicken Salad",   // 2-100 chars\n  "ingredients": [{"name": "chicken breast", "quantity": "200g"}],  // min 1\n  "directions": ["Grill chicken", "Toss with greens"],\n  "category": "High Protein",\n  "servings": 2,\n  "calories": 350, "protein_g": 40, "carbs_g": 15, "fat_g": 12\n}')

doc.add_page_break()

# ─── 8. Meal Plans APIs ───
add_colored_heading('8. Meal Plans APIs')
doc.add_paragraph('Source: sprint4.py — Weekly meal planning with breakfast/lunch/dinner slots.')
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/v1/meal-plan', 'Add meal to weekly plan (meal_id, date, slot)', 'Yes'],
    ['GET', '/api/v1/meal-plan', 'Get meal plan (?date=2026-04-02 or ?startDate&endDate)', 'Yes'],
    ['DELETE', '/api/v1/meal-plan/{id}', 'Remove meal from plan', 'Yes'],
])

doc.add_page_break()

# ─── 9. Trackers APIs ───
add_colored_heading('9. Trackers APIs')
doc.add_paragraph('Source: sprint3.py — Health & wellness trackers for water, sleep, walking, MET, and happiness.')

tracker_modules = [
    ('9.1 Water Tracker', [
        ['POST', '/api/v1/trackers/water', 'Log water intake', '{ glasses: 1 }'],
        ['GET', '/api/v1/trackers/water', 'Get water logs for today', 'Returns { total_glasses, goal, logs }'],
    ]),
    ('9.2 Sleep Tracker', [
        ['POST', '/api/v1/trackers/sleep', 'Log sleep data', '{ bedtime, wake_time, duration, quality (1-5) }'],
        ['GET', '/api/v1/trackers/sleep', 'Get sleep logs', 'Returns last 7 days of sleep data'],
    ]),
    ('9.3 Walking Tracker', [
        ['POST', '/api/v1/trackers/walking', 'Log walking/step data', '{ steps, distance, calories, duration }'],
        ['GET', '/api/v1/trackers/walking', 'Get walking logs', 'Returns step history'],
    ]),
    ('9.4 MET Activity Tracker', [
        ['POST', '/api/v1/trackers/met', 'Log MET activity', '{ activity_type, met_value, duration, met_minutes }'],
        ['GET', '/api/v1/trackers/met', 'Get MET logs', 'Activity types: Walking/Cycling/Swimming/Running/Yoga/Strength'],
    ]),
    ('9.5 Happiness / Mood Tracker', [
        ['POST', '/api/v1/trackers/happiness', 'Log happiness level (1-5)', '{ level: 4, note: "Great day!" }'],
        ['GET', '/api/v1/trackers/happiness', 'Get happiness log history', 'Returns daily mood entries'],
        ['POST', '/api/v1/happiness', 'Log happiness (alt endpoint)', '{ level: 4, factors: [...] }'],
        ['GET', '/api/v1/happiness/today', 'Get today\'s happiness entry', 'Single entry or null'],
        ['GET', '/api/v1/happiness/history', 'Get mood history with trends', 'Includes 30-day trend analysis'],
        ['GET', '/api/v1/happiness/quote', 'Get a motivational quote', 'Random from 30 seeded quotes'],
    ]),
    ('9.6 Tracker Aggregation', [
        ['GET', '/api/v1/trackers/summary', 'Daily tracker summary (all trackers)', 'Aggregated water, sleep, steps, MET, mood'],
        ['GET', '/api/v1/trackers/timeline', 'Daily activity timeline', 'Chronological view of all logged activities'],
        ['GET', '/api/v1/progress/overview', 'Full progress overview with streaks', 'Comprehensive health metrics'],
    ]),
]

for title, endpoints in tracker_modules:
    doc.add_heading(title, level=2)
    add_table(['Method', 'Endpoint', 'Description', 'Details'], endpoints)

doc.add_page_break()

# ─── 10. Journal APIs ───
add_colored_heading('10. Journal APIs')
doc.add_paragraph('Source: sprint3.py — Daily journaling for wellness reflections.')
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/v1/journal', 'Create journal entry', 'Yes'],
    ['GET', '/api/v1/journal', 'List journal entries (paginated)', 'Yes'],
    ['GET', '/api/v1/journal/{id}', 'Get single journal entry', 'Yes'],
    ['PUT', '/api/v1/journal/{id}', 'Update journal entry', 'Yes'],
    ['DELETE', '/api/v1/journal/{id}', 'Delete journal entry', 'Yes'],
    ['POST', '/api/v1/journal/like', 'Toggle like on a journal entry', 'Yes'],
])

doc.add_page_break()

# ─── 11. Goals & Wellness ───
add_colored_heading('11. Goals & Wellness Program APIs')
doc.add_paragraph('Source: sprint3.py, sprint2.py')

doc.add_heading('11.1 Goals', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/goals', 'Get user goals (life goals, wellness goals, etc.)'],
    ['GET', '/api/v1/goals/progress', 'Get goal progress cards (walking, nutrition, activity, hydration) with milestone badges'],
])

doc.add_heading('11.2 Wellness Programs', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/wellness-programs', 'List all active wellness programs (4 programs)'],
    ['GET', '/api/wellness-programs/{id}', 'Get program detail'],
    ['POST', '/api/v1/wellness-programs/{id}/enroll', 'Enroll in a wellness program'],
    ['POST', '/api/v1/wellness-programs/checkin', 'Daily check-in for active program'],
    ['GET', '/api/v1/wellness-programs/progress/{enrollment_id}', 'Get enrollment progress'],
    ['GET', '/api/v1/wellness-programs/active', 'Get user\'s active program enrollment'],
])

doc.add_heading('11.3 Reports', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['POST', '/api/v1/reports/generate', 'Generate 30-day progress report across all trackers'],
])

doc.add_page_break()

# ─── 12. Social Feed APIs ───
add_colored_heading('12. Social Feed APIs')
doc.add_paragraph('Sources: server.py (basic feed), sprint4.py (enhanced feed with pagination, ownership, media)')

doc.add_heading('12.1 Posts', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/v1/feed', 'Create a feed post (text + optional image_url)', 'Yes'],
    ['GET', '/api/v1/feed', 'Get paginated feed (?page=1&limit=20)', 'Yes'],
    ['GET', '/api/v1/feed/{id}', 'Get post detail', 'Yes'],
    ['PUT', '/api/v1/feed/{id}', 'Update own post', 'Yes (owner)'],
    ['DELETE', '/api/v1/feed/{id}', 'Delete own post', 'Yes (owner)'],
])

doc.add_heading('12.2 Interactions', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/v1/post/like/{postId}', 'Toggle like on a post', 'Yes'],
    ['GET', '/api/v1/post/likes/{postId}', 'Get list of users who liked', 'Yes'],
    ['POST', '/api/v1/post/comment/{postId}', 'Add comment to post', 'Yes'],
    ['GET', '/api/v1/post/comments/{postId}', 'Get comments for a post', 'Yes'],
    ['PUT', '/api/v1/post/{postId}/comment/{commentId}', 'Edit own comment', 'Yes (owner)'],
    ['DELETE', '/api/v1/post/{postId}/comment/{commentId}', 'Delete own comment', 'Yes (owner)'],
])

doc.add_page_break()

# ─── 13. AI Chat APIs ───
add_colored_heading('13. AI Wellness Coach (Chat) APIs')
doc.add_paragraph('Source: server.py — GPT-4.1-mini powered wellness chatbot with user-context awareness.')
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/chat', 'Send message to AI coach, get response', 'Yes'],
    ['GET', '/api/chat/history', 'Get full chat history (up to 100 messages)', 'Yes'],
])
doc.add_paragraph('Chat Request:')
add_code_block('{ "message": "What should I eat for a high-protein lunch?" }')
doc.add_paragraph('Chat Response:')
add_code_block('{\n  "user_message": { "id": "...", "role": "user", "content": "..." },\n  "ai_message": { "id": "...", "role": "assistant", "content": "Here are some great options..." }\n}')
doc.add_paragraph('The AI system prompt includes user profile data (goals, dietary preferences, allergies, weight, activity level) for personalized responses.')

doc.add_page_break()

# ─── 14. Wearable APIs ───
add_colored_heading('14. Wearable Device APIs')
doc.add_paragraph('Source: wearable.py — Connect, sync, and manage wearable health devices.')
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['GET', '/api/v1/wearables/providers', 'List supported providers (5 providers)', 'Yes'],
    ['GET', '/api/v1/wearables/connected', 'Get user\'s connected devices', 'Yes'],
    ['POST', '/api/v1/wearables/connect', 'Connect a wearable device', 'Yes'],
    ['DELETE', '/api/v1/wearables/disconnect/{provider}', 'Disconnect a device', 'Yes'],
    ['POST', '/api/v1/wearables/sync', 'Sync health data from device', 'Yes'],
    ['POST', '/api/v1/wearables/data', 'Submit individual data point', 'Yes'],
    ['GET', '/api/v1/wearables/data', 'Get wearable data history', 'Yes'],
    ['GET', '/api/v1/wearables/summary', 'Get aggregated health summary', 'Yes'],
])
doc.add_paragraph('Supported Providers:')
add_table(['Provider ID', 'Name', 'Platforms'], [
    ['apple_health', 'Apple Health', 'iOS'],
    ['google_fit', 'Google Fit', 'Android'],
    ['fitbit', 'Fitbit', 'iOS, Android'],
    ['samsung_health', 'Samsung Health', 'Android'],
    ['garmin', 'Garmin Connect', 'iOS, Android'],
])
doc.add_paragraph('Sync Data Types: steps, heart_rate, calories, sleep, distance, active_minutes')

doc.add_page_break()

# ─── 15. Payment APIs ───
add_colored_heading('15. Payment & Subscription APIs')
doc.add_paragraph('Sources: payment.py (Stripe), sprint5.py (subscription management)')

doc.add_heading('15.1 Stripe Checkout', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['GET', '/api/v1/payment/config', 'Get Stripe publishable key', 'Yes'],
    ['POST', '/api/v1/payment/create-checkout', 'Create Stripe Checkout session', 'Yes'],
    ['POST', '/api/v1/payment/confirm', 'Confirm payment after checkout', 'Yes'],
    ['GET', '/api/v1/payment/history', 'Get payment transaction history', 'Yes'],
])
doc.add_paragraph('Create Checkout Request:')
add_code_block('{\n  "plan_id": "664f...",\n  "success_url": "https://your-app.com/payment/success",\n  "cancel_url": "https://your-app.com/payment/cancel"\n}')

doc.add_heading('15.2 Subscription Management', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['GET', '/api/v1/subscription/plans', 'List all subscription plans (3 tiers)', 'Yes'],
    ['POST', '/api/v1/subscription', 'Purchase subscription', 'Yes'],
    ['GET', '/api/v1/subscription', 'Get current subscription status', 'Yes'],
    ['PUT', '/api/v1/subscription/cancel', 'Cancel subscription', 'Yes'],
    ['GET', '/api/v1/subscription/transactions', 'Get transaction history (paginated)', 'Yes'],
])

doc.add_page_break()

# ─── 16. Notification APIs ───
add_colored_heading('16. Notifications & Push APIs')
doc.add_paragraph('Sources: sprint5.py (in-app notifications), push_service.py (Expo Push), server.py (admin push routes)')

doc.add_heading('16.1 In-App Notifications', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/v1/notifications/register', 'Register device push token', 'Yes'],
    ['GET', '/api/v1/notifications', 'Get user notifications (paginated, ?unreadOnly=true)', 'Yes'],
    ['PUT', '/api/v1/notifications/{id}/read', 'Mark notification as read', 'Yes'],
    ['PUT', '/api/v1/notifications/read-all', 'Mark all notifications as read', 'Yes'],
    ['DELETE', '/api/v1/notifications/{id}', 'Delete a notification', 'Yes'],
    ['GET', '/api/v1/notifications/preferences', 'Get notification preferences', 'Yes'],
    ['PUT', '/api/v1/notifications/preferences', 'Update notification preferences', 'Yes'],
])

doc.add_heading('16.2 Admin Push Operations', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/v1/push/broadcast', 'Broadcast push to all devices', 'Admin'],
    ['POST', '/api/v1/push/user', 'Send push to specific user', 'Admin'],
    ['POST', '/api/v1/push/happiness-reminder', 'Trigger happiness check-in reminder', 'Admin'],
    ['POST', '/api/v1/notifications/broadcast', 'Broadcast in-app notification', 'Admin'],
])

doc.add_page_break()

# ─── 17. Restaurant APIs ───
add_colored_heading('17. Restaurant APIs')
doc.add_paragraph('Source: sprint2.py — Restaurant discovery, search, reviews, favorites, and claiming.')
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['GET', '/api/restaurants', 'List restaurants (paginated, sorted)', 'No'],
    ['GET', '/api/restaurants/search', 'Advanced search (cuisine, rating, distance, BO verified)', 'No'],
    ['GET', '/api/restaurants/nearby', 'Find nearby restaurants (lat/lng/radius)', 'No'],
    ['GET', '/api/restaurants/{id}', 'Get restaurant detail', 'No'],
    ['GET', '/api/restaurants/menu/{id}', 'Get restaurant menu items', 'No'],
    ['GET', '/api/restaurants/favorites', 'Get user\'s favorite restaurants', 'Yes'],
    ['POST', '/api/restaurants/like/{id}', 'Toggle restaurant favorite', 'Yes'],
    ['POST', '/api/restaurants/rating', 'Submit rating (1-5)', 'Yes'],
    ['POST', '/api/restaurants/reviews', 'Submit a review', 'Yes'],
    ['GET', '/api/restaurants/reviews/{id}', 'Get restaurant reviews (paginated)', 'No'],
    ['GET', '/api/restaurants/boVerified', 'Check BO verification status', 'No'],
    ['GET', '/api/restaurants/boPartner', 'Check BO partner status', 'No'],
    ['POST', '/api/v1/restaurants/claims', 'Submit restaurant ownership claim', 'Yes'],
    ['GET', '/api/v1/restaurants/claims/mine', 'Get user\'s submitted claims', 'Yes'],
])
doc.add_paragraph('Search Parameters:')
add_code_block('GET /api/restaurants/search?q=sushi&cuisine=Japanese&min_rating=4&max_distance=10&bo_verified=true&lat=40.73&lng=-73.99')

doc.add_page_break()

# ─── 18. Support & FAQ APIs ───
add_colored_heading('18. Support Tickets & FAQ APIs')
doc.add_paragraph('Source: sprint6.py')

doc.add_heading('18.1 Support Tickets', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/v1/ticket', 'Create support ticket', 'Yes'],
    ['GET', '/api/v1/tickets', 'List user\'s tickets', 'Yes'],
    ['GET', '/api/v1/tickets/{id}', 'Get ticket detail with messages', 'Yes'],
    ['PUT', '/api/v1/tickets/{id}', 'Update ticket', 'Yes'],
    ['POST', '/api/v1/ticket/message', 'Add message to ticket thread', 'Yes'],
    ['DELETE', '/api/v1/tickets/{id}', 'Delete/close ticket', 'Yes'],
    ['GET', '/api/v1/tickets/allmessages', 'Get all messages across tickets', 'Yes'],
])

doc.add_heading('18.2 FAQ', level=2)
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['GET', '/api/v1/faqs', 'List all FAQ entries (20 seeded)', 'No'],
    ['GET', '/api/v1/faq/{id}', 'Get single FAQ detail', 'No'],
])

doc.add_page_break()

# ─── 19. Referrals & Account ───
add_colored_heading('19. Referrals & Account Management APIs')
doc.add_paragraph('Source: sprint6.py')
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['POST', '/api/v1/referrals/generate', 'Generate a unique referral code', 'Yes'],
    ['GET', '/api/v1/referrals', 'Get referral code and stats', 'Yes'],
    ['GET', '/api/v1/legal/terms', 'Get Terms of Service text', 'No'],
    ['GET', '/api/v1/legal/privacy', 'Get Privacy Policy text', 'No'],
    ['GET', '/api/v1/app/version', 'Get current app version info', 'No'],
    ['POST', '/api/v1/account/delete-request', 'Request account deletion (30-day grace)', 'Yes'],
    ['POST', '/api/v1/account/reactivate', 'Reactivate a deletion-pending account', 'Yes'],
])

doc.add_page_break()

# ─── 20. Admin Panel APIs ───
add_colored_heading('20. Admin Panel APIs')
doc.add_paragraph('Sources: sprint7.py, sprint8.py, sprint9.py, admin_panel.py — Web-based admin panel with 2FA.')

doc.add_heading('20.1 Admin Authentication', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['POST', '/api/v1/admin/login', 'Admin login (returns session + 2FA challenge)'],
    ['POST', '/api/v1/admin/verify-2fa', 'Verify 2FA OTP code'],
    ['POST', '/api/v1/auth/demo-login', 'Demo admin login (dev only)'],
    ['GET', '/api/admin-panel', 'Serve admin panel HTML (web dashboard)'],
])

doc.add_heading('20.2 Dashboard & Analytics', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/admin/dashboard', 'KPI stats (users, revenue, meals, posts, tickets)'],
    ['GET', '/api/v1/admin/notifications/analytics', 'Push notification analytics'],
    ['GET', '/api/v1/admin/subscription-plans/analytics', 'Subscription analytics'],
])

doc.add_heading('20.3 User Management', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/admin/users', 'List all users (paginated, searchable)'],
    ['GET', '/api/v1/admin/users/subscribed', 'List subscribed/pro users'],
    ['GET', '/api/v1/admin/users/{id}', 'Get user detail'],
    ['GET', '/api/v1/admin/user/all-data/{id}', 'Get ALL user data (meals, trackers, posts, etc.)'],
    ['POST', '/api/v1/admin/users/changeAction/{id}', 'Suspend / activate / delete user'],
    ['POST', '/api/v1/admin/users/impersonate/{id}', 'Generate impersonation token'],
    ['POST', '/api/v1/admin/users/create-admin', 'Create new admin user'],
    ['GET', '/api/v1/admin/team', 'List admin team members'],
])

doc.add_heading('20.4 Content Management', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/admin/meal', 'List meals (paginated, filtered)'],
    ['POST', '/api/v1/admin/meal', 'Create a new meal'],
    ['PUT', '/api/v1/admin/meal/{id}', 'Update meal'],
    ['DELETE', '/api/v1/admin/meal/{id}', 'Delete meal'],
    ['PUT', '/api/v1/admin/meal/{id}/approve', 'Approve a meal'],
    ['PUT', '/api/v1/admin/meal/{id}/reject', 'Reject a meal'],
    ['GET', '/api/v1/admin/ingredients/suggest', 'AI ingredient suggestions'],
    ['GET', '/api/v1/admin/posts', 'List feed posts'],
    ['POST', '/api/v1/admin/post', 'Create admin post'],
    ['PUT', '/api/v1/admin/post/{id}', 'Update post'],
    ['DELETE', '/api/v1/admin/post/{id}', 'Delete post'],
])

doc.add_heading('20.5 Quotes Management', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/admin/quotes', 'List all quotes'],
    ['POST', '/api/v1/admin/quotes', 'Create a new quote'],
    ['PUT', '/api/v1/admin/quotes/{id}', 'Update quote'],
    ['DELETE', '/api/v1/admin/quotes/{id}', 'Delete quote'],
    ['POST', '/api/v1/admin/select/quotes/{id}', 'Select quote as daily'],
    ['GET', '/api/v1/admin/selected', 'Get currently selected quote'],
    ['GET', '/api/v1/quotes/today', 'Get today\'s quote (public)'],
])

doc.add_heading('20.6 Subscription Plans', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/admin/subscription-plans', 'List subscription plans'],
    ['POST', '/api/v1/admin/subscription-plan', 'Create plan'],
    ['PUT', '/api/v1/admin/subscription-plan/{id}', 'Update plan'],
    ['DELETE', '/api/v1/admin/subscription-plan/{id}', 'Delete plan'],
])

doc.add_heading('20.7 Claims & Tickets', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/admin/claims', 'List all restaurant claims'],
    ['PUT', '/api/v1/admin/claims/{id}/approve', 'Approve a claim'],
    ['PUT', '/api/v1/admin/claims/{id}/reject', 'Reject a claim'],
    ['GET', '/api/v1/admin/tickets', 'List all support tickets'],
    ['GET', '/api/v1/admin/tickets/{id}', 'Get ticket detail'],
    ['PUT', '/api/v1/admin/ticket/change_status/{id}', 'Change ticket status'],
    ['PUT', '/api/v1/admin/tickets/{id}', 'Update ticket / add response'],
    ['POST', '/api/v1/admin/ticket/message', 'Add admin message to ticket'],
    ['DELETE', '/api/v1/admin/tickets/{id}', 'Delete ticket'],
    ['POST', '/api/v1/admin/tickets/report', 'Generate ticket analytics report'],
])

doc.add_heading('20.8 Admin Settings & Sessions', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['PUT', '/api/v1/admin/profile', 'Update admin profile'],
    ['GET', '/api/v1/admin/sessions', 'List active admin sessions'],
    ['DELETE', '/api/v1/admin/sessions/{id}', 'Revoke an admin session'],
])

doc.add_heading('20.9 Restaurant & Distributor Management', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/admin/restaurants', 'List restaurants'],
    ['POST', '/api/v1/admin/restaurants', 'Create restaurant'],
    ['PUT', '/api/v1/admin/restaurants/{id}', 'Update restaurant'],
    ['DELETE', '/api/v1/admin/restaurants/{id}', 'Delete restaurant'],
    ['GET', '/api/v1/admin/distributors', 'List distributors'],
    ['POST', '/api/v1/admin/distributors', 'Create distributor'],
    ['PUT', '/api/v1/admin/distributors/{id}', 'Update distributor'],
    ['DELETE', '/api/v1/admin/distributors/{id}', 'Delete distributor'],
])

doc.add_heading('20.10 Admin FAQ Management', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['GET', '/api/v1/admin/faqs', 'List all FAQs (admin view)'],
    ['POST', '/api/v1/admin/faq', 'Create FAQ entry'],
    ['PUT', '/api/v1/admin/faq/{id}', 'Update FAQ'],
    ['DELETE', '/api/v1/admin/faq/{id}', 'Delete FAQ'],
])

doc.add_heading('20.11 Admin Notification Management', level=2)
add_table(['Method', 'Endpoint', 'Description'], [
    ['POST', '/api/v1/admin/notifications/broadcast', 'Broadcast notification to all users'],
    ['GET', '/api/v1/admin/notifications/history', 'Notification broadcast history'],
    ['GET', '/api/v1/admin/notifications/analytics', 'Notification delivery analytics'],
])

doc.add_page_break()

# ─── 21. Admin Enterprise APIs ───
add_colored_heading('21. Admin Enterprise APIs (AI Analytics, 360 View)')
doc.add_paragraph('Source: admin_enterprise.py — AI-powered analytics and comprehensive user insights.')

doc.add_heading('GET /api/v1/admin/ai-analytics', level=2)
add_info('Description', 'AI-generated analytics dashboard with GPT-4.1-mini powered insights')
add_info('Auth', 'Admin (session token)')
doc.add_paragraph('Response includes: user growth trends, engagement metrics, health outcome insights, revenue analysis, retention analysis, and AI-generated recommendations.')

doc.add_heading('POST /api/v1/admin/ai-recipe-info', level=2)
add_info('Description', 'AI-powered nutritional information generation for meals')
add_info('Auth', 'Admin')
doc.add_paragraph('Request Body:')
add_code_block('{ "meal_name": "Grilled Chicken Caesar Salad" }')
doc.add_paragraph('Response: Auto-generated calories, protein, carbs, fat, ingredients, directions with "approx" labels.')

doc.add_heading('GET /api/v1/admin/user/360/{user_id}', level=2)
add_info('Description', '360-degree user profile view with all activity data')
add_info('Auth', 'Admin')
doc.add_paragraph('Response: Complete user data including health timeline, workout history, meal logs, social activity, subscription history, support tickets, wearable data, and badges.')

doc.add_page_break()

# ─── 22. Utility APIs ───
add_colored_heading('22. Utility & Health Check APIs')
add_table(['Method', 'Endpoint', 'Description', 'Auth'], [
    ['GET', '/api/v1/health', 'Health check (DB status, collection count, version)', 'No'],
    ['GET', '/api/meal-categories', 'List meal categories (9 categories)', 'No'],
    ['GET', '/api/meal-featured', 'Get random featured meal', 'No'],
    ['POST', '/api/v1/upload', 'Upload media file to Cloudinary', 'Yes'],
    ['GET', '/api/v1/predictions', 'AI-powered health predictions', 'Yes'],
    ['GET', '/api/download/project-docs', 'Download project documentation (.docx)', 'No'],
    ['GET', '/api/v1/badges', 'List all badges (12 badges, 4 categories)', 'Yes'],
    ['GET', '/api/v1/badges/check', 'Check earned badges', 'Yes'],
    ['GET', '/api/v1/badges/progress', 'Get badge progress', 'Yes'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# PART B: EXTERNAL API INTEGRATIONS
# ════════════════════════════════════════════════════════════
ph = doc.add_paragraph()
ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
ph.space_before = Pt(40)
rh = ph.add_run('PART B')
rh.font.size = Pt(36)
rh.font.bold = True
rh.font.color.rgb = GREEN
ph2 = doc.add_paragraph()
ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rh2 = ph2.add_run('External API & Third-Party Integrations')
rh2.font.size = Pt(20)
rh2.font.color.rgb = DARK
doc.add_page_break()

# ─── 23. OpenAI / Emergent ───
add_colored_heading('23. OpenAI (GPT-4.1-mini) via Emergent Integrations')

add_table(['Property', 'Value'], [
    ['Provider', 'OpenAI'],
    ['Model', 'gpt-4.1-mini'],
    ['Access Library', 'emergentintegrations (Python)'],
    ['API Key Variable', 'EMERGENT_LLM_KEY'],
    ['Auth Mechanism', 'Universal Emergent LLM Key (works across OpenAI, Anthropic, Google)'],
    ['Rate Limits', 'Per Emergent plan (typically 60 RPM for dev tier)'],
    ['Pricing', 'Covered by Emergent universal key subscription'],
])

doc.add_heading('Usage in BO App', level=2)
add_table(['Feature', 'File', 'How Used'], [
    ['AI Wellness Coach', 'server.py', 'Chat endpoint — personalized health advice with user context (goals, diet, allergies, weight)'],
    ['AI Analytics Dashboard', 'admin_enterprise.py', 'Generates platform health insights, trends, and recommendations for admin panel'],
    ['AI Recipe Generator', 'admin_enterprise.py', 'Auto-generates nutritional info (calories, macros, ingredients, directions) for new meals'],
])

doc.add_heading('Integration Code Pattern', level=2)
add_code_block('''from emergentintegrations.llm.chat import LlmChat, UserMessage

chat = LlmChat(
    api_key=os.environ.get("EMERGENT_LLM_KEY"),
    session_id="unique-session-id",
    system_message="You are a wellness AI coach..."
)
chat.with_model("openai", "gpt-4.1-mini")

user_message = UserMessage(text="What should I eat today?")
response = await chat.send_message(user_message)
# response is a string containing the AI's reply''')

doc.add_heading('External API Called', level=2)
add_table(['Property', 'Value'], [
    ['Base URL', 'https://api.openai.com/v1/chat/completions (proxied via Emergent)'],
    ['Method', 'POST'],
    ['Content-Type', 'application/json'],
    ['Model Used', 'gpt-4.1-mini'],
    ['Max Tokens', 'Default (varies by Emergent config)'],
    ['Temperature', 'Default 0.7'],
])

doc.add_page_break()

# ─── 24. Stripe ───
add_colored_heading('24. Stripe Payment Gateway')

add_table(['Property', 'Value'], [
    ['Provider', 'Stripe'],
    ['SDK', 'stripe (Python)'],
    ['Version', 'Latest via pip'],
    ['API Key Variables', 'STRIPE_SECRET_KEY, STRIPE_PUBLISHABLE_KEY'],
    ['Mode', 'Test mode (sk_test_*, pk_test_*)'],
    ['Dashboard', 'https://dashboard.stripe.com'],
    ['Documentation', 'https://stripe.com/docs/api'],
])

doc.add_heading('Stripe APIs Used by BO', level=2)
add_table(['Stripe API', 'Method', 'BO Endpoint', 'Purpose'], [
    ['Checkout Sessions', 'POST /v1/checkout/sessions', '/api/v1/payment/create-checkout', 'Create a hosted checkout page for subscription purchase'],
    ['Session Retrieve', 'GET /v1/checkout/sessions/{id}', '/api/v1/payment/confirm', 'Verify payment completion and retrieve session details'],
])

doc.add_heading('Integration Flow', level=2)
flow_steps = [
    '1. User selects a subscription plan in the app',
    '2. App calls POST /api/v1/payment/create-checkout with plan_id',
    '3. Backend creates a Stripe Checkout Session with line items, customer email, and success/cancel URLs',
    '4. Backend returns session URL to the app',
    '5. App opens the Stripe-hosted checkout page in a WebView',
    '6. After payment, Stripe redirects to success URL with ?session_id=...',
    '7. App calls POST /api/v1/payment/confirm with the session_id',
    '8. Backend retrieves the session from Stripe, verifies payment status, and activates subscription',
]
for step in flow_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('Required API Keys', level=2)
add_table(['Key', 'Purpose', 'Where to Get'], [
    ['STRIPE_SECRET_KEY', 'Server-side API operations', 'Stripe Dashboard > Developers > API Keys'],
    ['STRIPE_PUBLISHABLE_KEY', 'Client-side checkout initialization', 'Stripe Dashboard > Developers > API Keys'],
])

doc.add_heading('Webhook (Future)', level=2)
doc.add_paragraph('Stripe webhooks are not yet configured. For production, a webhook endpoint should be added to handle events like checkout.session.completed, customer.subscription.updated, and invoice.payment_failed.')

doc.add_page_break()

# ─── 25. Cloudinary ───
add_colored_heading('25. Cloudinary Image CDN')

add_table(['Property', 'Value'], [
    ['Provider', 'Cloudinary'],
    ['SDK', 'cloudinary (Python)'],
    ['Purpose', 'Image upload, storage, and transformation for profile photos and feed media'],
    ['API Key Variables', 'CLOUDINARY_CLOUD_NAME, CLOUDINARY_API_KEY, CLOUDINARY_API_SECRET'],
    ['Dashboard', 'https://console.cloudinary.com'],
    ['Documentation', 'https://cloudinary.com/documentation'],
])

doc.add_heading('Cloudinary APIs Used by BO', level=2)
add_table(['API', 'Method', 'BO Usage'], [
    ['Upload API', 'POST /v1_1/{cloud}/image/upload', 'Upload profile photos and feed post images'],
    ['Transform URLs', 'GET /image/upload/w_400,h_400,c_fill/', 'On-the-fly image optimization via URL transforms'],
])

doc.add_heading('Upload Integration (sprint5.py)', level=2)
add_code_block('''import cloudinary
import cloudinary.uploader

cloudinary.config(
    cloud_name=os.environ.get("CLOUDINARY_CLOUD_NAME"),
    api_key=os.environ.get("CLOUDINARY_API_KEY"),
    api_secret=os.environ.get("CLOUDINARY_API_SECRET"),
)

# Upload endpoint: POST /api/v1/upload
result = cloudinary.uploader.upload(
    file_content,
    folder="bo-wellness",
    resource_type="image",
)
# Returns: { "secure_url": "https://res.cloudinary.com/..." }''')

doc.add_heading('Image Optimization (middleware.py)', level=2)
doc.add_paragraph('The middleware automatically optimizes Cloudinary URLs with quality and size transforms:')
add_table(['Transform', 'Parameters', 'Use Case'], [
    ['thumbnail', 'w_150,h_150,c_fill,q_60,f_auto', 'Avatar thumbnails, list views'],
    ['medium', 'w_400,h_400,c_fill,q_70,f_auto', 'Feed images, meal photos'],
    ['feed', 'w_600,h_400,c_fill,q_75,f_auto', 'Social feed cards'],
    ['large', 'w_800,h_600,c_fill,q_80,f_auto', 'Detail view images'],
])

doc.add_heading('Required API Keys', level=2)
add_table(['Key', 'Purpose', 'Where to Get'], [
    ['CLOUDINARY_CLOUD_NAME', 'Identifies your Cloudinary account', 'Cloudinary Dashboard > Account Details'],
    ['CLOUDINARY_API_KEY', 'API authentication', 'Cloudinary Dashboard > Account Details'],
    ['CLOUDINARY_API_SECRET', 'API secret for server-side operations', 'Cloudinary Dashboard > Account Details'],
])

doc.add_page_break()

# ─── 26. Expo Push ───
add_colored_heading('26. Expo Push Notification Service')

add_table(['Property', 'Value'], [
    ['Provider', 'Expo Push Service'],
    ['Endpoint', 'https://exp.host/--/api/v2/push/send'],
    ['SDK', 'httpx (HTTP client, no SDK needed)'],
    ['Auth', 'No API key required (token-based)'],
    ['Documentation', 'https://docs.expo.dev/push-notifications/overview/'],
    ['Batch Limit', '100 messages per API call'],
])

doc.add_heading('How It Works', level=2)
flow = [
    '1. User opens app, device registers for push → gets ExponentPushToken[xxx]',
    '2. App sends token to POST /api/v1/notifications/register',
    '3. Backend stores token in push_tokens collection linked to user_id',
    '4. When a notification is triggered, backend calls Expo Push API',
    '5. Expo routes the notification to Apple APNs (iOS) or Google FCM (Android)',
    '6. User receives native push notification on their device',
]
for step in flow:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('API Request Format', level=2)
add_code_block('''POST https://exp.host/--/api/v2/push/send
Content-Type: application/json

{
  "to": "ExponentPushToken[xxxxxx]",
  "title": "BO | Meal Reminder",
  "body": "Time for your healthy lunch!",
  "sound": "default",
  "channelId": "default",
  "subtitle": "BO Wellness",
  "data": { "deepLink": "/home", "type": "meal_reminder" }
}''')

doc.add_heading('Bulk Sending', level=2)
doc.add_paragraph('For broadcasts, the service sends messages in batches of 100 (Expo API limit). The push_service.py handles batching automatically.')

doc.add_heading('BO Branding', level=2)
doc.add_paragraph('All push notification titles are automatically prefixed with "BO | " if not already starting with "BO". The subtitle is always set to "BO Wellness".')

doc.add_heading('Notification Types Sent', level=2)
add_table(['Type', 'Trigger', 'Title', 'Body Example'], [
    ['Meal Reminder', 'Scheduled (3x daily)', 'BO | Meal Reminder', 'Time for your healthy lunch!'],
    ['Water Reminder', 'Scheduled (hourly)', 'BO | Hydration Check', 'Remember to drink water!'],
    ['Happiness Check-in', 'Daily (if not logged)', 'BO | How are you feeling?', 'Take a moment to check in...'],
    ['Wellness Tip', 'Admin broadcast', 'BO | Wellness Tip', 'Custom message from admin'],
    ['Admin Broadcast', 'Manual', 'Custom', 'Custom message'],
])

doc.add_page_break()

# ─── 27. MongoDB ───
add_colored_heading('27. MongoDB (Motor Async Driver)')

add_table(['Property', 'Value'], [
    ['Provider', 'MongoDB (local in dev, Atlas in prod)'],
    ['Driver', 'Motor (async Python driver for MongoDB)'],
    ['Connection Variable', 'MONGO_URL'],
    ['Database Name Variable', 'DB_NAME'],
    ['Index Strategy', 'Created on app startup via seed functions'],
])

doc.add_heading('Collections (29 total)', level=2)
add_table(['Collection', 'Purpose', 'Key Indexes'], [
    ['users', 'User accounts, profiles, preferences', 'email (unique)'],
    ['meals', 'Curated meal database (66 meals)', 'plan_id, category'],
    ['restaurants', 'Restaurant listings (60)', 'location (2dsphere), is_active'],
    ['menu_items', 'Restaurant menu items', 'restaurant_id'],
    ['feed_posts / posts', 'Community feed posts', 'created_at'],
    ['comments', 'Post comments', 'post_id'],
    ['nutrition_logs', 'Nutrition tracking entries', 'user_id + date'],
    ['water_logs', 'Water intake records', 'user_id + date'],
    ['sleep_logs', 'Sleep tracking records', 'user_id + date'],
    ['walking_logs', 'Step/walking records', 'user_id + date'],
    ['met_logs', 'MET activity records', 'user_id + date'],
    ['happiness_logs', 'Mood tracking records', 'user_id + date'],
    ['workouts', 'Workout records', 'user_id'],
    ['journals', 'Journal entries', 'user_id + created_at'],
    ['chat_messages', 'AI chat history', 'user_id + created_at'],
    ['wellness_programs', 'Program definitions (4)', '-'],
    ['wellness_enrollments', 'User program enrollments', 'user_id'],
    ['badges', 'Badge definitions (12)', '-'],
    ['user_badges', 'Earned badges', 'user_id'],
    ['notifications', 'In-app notifications', 'user_id + created_at'],
    ['push_tokens', 'Device push tokens', 'user_id'],
    ['subscription_plans', 'Plan definitions (3)', '-'],
    ['user_subscriptions', 'Active subscriptions', 'user_id'],
    ['stripe_sessions', 'Stripe checkout records', 'session_id'],
    ['restaurant_claims', 'Ownership claims', 'user_id + restaurant_id'],
    ['restaurant_favorites', 'User favorites', 'user_id + restaurant_id (unique)'],
    ['restaurant_ratings', 'User ratings', 'user_id + restaurant_id (unique)'],
    ['restaurant_reviews', 'User reviews', 'restaurant_id + created_at'],
    ['wearable_connections', 'Connected devices', 'user_id + provider'],
    ['wearable_data', 'Synced health data', 'user_id + data_type + recorded_at'],
    ['admin_quotes', 'Daily motivational quotes (30)', '-'],
    ['diet_plans', 'Diet plan definitions (6)', '-'],
    ['faqs', 'FAQ entries (20)', '-'],
    ['reset_codes', 'Password reset codes', 'email'],
    ['meal_categories', 'Meal categories (9)', '-'],
    ['weight_logs', 'Weight tracking', 'user_id'],
    ['tickets', 'Support tickets', 'user_id'],
    ['ticket_messages', 'Ticket thread messages', 'ticket_id'],
    ['referrals', 'Referral codes & tracking', 'user_id'],
])

doc.add_page_break()

# ─── 28. Wearable Providers ───
add_colored_heading('28. Wearable Provider APIs')
doc.add_paragraph('The app supports 5 wearable providers. Currently, the integration uses a REST-based data sync pattern where the mobile app collects data from native health APIs and sends it to the backend.')

add_table(['Provider', 'Native API', 'Platforms', 'Data Types', 'Integration Status'], [
    ['Apple Health', 'HealthKit (iOS SDK)', 'iOS only', 'steps, HR, calories, sleep, distance, active_min', 'REST sync via app'],
    ['Google Fit', 'Google Fit REST API', 'Android', 'steps, HR, calories, sleep, distance, active_min', 'REST sync via app'],
    ['Fitbit', 'Fitbit Web API (OAuth2)', 'iOS, Android', 'steps, HR, calories, sleep, distance, active_min', 'REST sync via app'],
    ['Samsung Health', 'Samsung Health SDK', 'Android', 'steps, HR, calories, sleep, distance, active_min', 'REST sync via app'],
    ['Garmin', 'Garmin Connect API', 'iOS, Android', 'steps, HR, calories, sleep, distance, active_min', 'REST sync via app'],
])

doc.add_heading('Future: Direct OAuth Integration', level=2)
doc.add_paragraph('For production, each provider requires OAuth2 setup:')
add_table(['Provider', 'OAuth URL', 'Required Credentials'], [
    ['Fitbit', 'https://dev.fitbit.com/', 'Client ID, Client Secret, OAuth2 redirect URI'],
    ['Google Fit', 'https://console.cloud.google.com/', 'OAuth2 Client ID, enabled Fitness API'],
    ['Garmin', 'https://developer.garmin.com/', 'Consumer Key, Consumer Secret, OAuth1.0a'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# PART C: APPENDICES
# ════════════════════════════════════════════════════════════
ph = doc.add_paragraph()
ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
ph.space_before = Pt(40)
rh = ph.add_run('PART C')
rh.font.size = Pt(36)
rh.font.bold = True
rh.font.color.rgb = GREEN
ph2 = doc.add_paragraph()
ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rh2 = ph2.add_run('Appendices')
rh2.font.size = Pt(20)
rh2.font.color.rgb = DARK
doc.add_page_break()

# ─── 29. Environment Variables ───
add_colored_heading('29. Environment Variables Reference')
add_table(['Variable', 'Required', 'Description', 'Example'], [
    ['MONGO_URL', 'Yes', 'MongoDB connection string', 'mongodb://localhost:27017'],
    ['DB_NAME', 'Yes', 'Database name', 'bo_wellness'],
    ['JWT_SECRET', 'Yes', 'Secret key for JWT token signing', 'bo-wellness-secret-key-2026'],
    ['ADMIN_EMAIL', 'Yes', 'Default admin email', 'admin@bo.com'],
    ['ADMIN_PASSWORD', 'Yes', 'Default admin password', 'BoAdmin2026!'],
    ['EMERGENT_LLM_KEY', 'Yes', 'Universal key for AI features (OpenAI, Anthropic, Google)', 'ek_xxxx...'],
    ['CLOUDINARY_CLOUD_NAME', 'For media', 'Cloudinary cloud name', 'my-cloud'],
    ['CLOUDINARY_API_KEY', 'For media', 'Cloudinary API key', '123456789'],
    ['CLOUDINARY_API_SECRET', 'For media', 'Cloudinary API secret', 'abc123...'],
    ['STRIPE_SECRET_KEY', 'For payments', 'Stripe secret key', 'sk_test_...'],
    ['STRIPE_PUBLISHABLE_KEY', 'For payments', 'Stripe publishable key', 'pk_test_...'],
])

doc.add_page_break()

# ─── 30. Error Codes ───
add_colored_heading('30. Error Code Reference')
add_table(['HTTP Code', 'Meaning', 'Common Causes'], [
    ['200', 'OK', 'Request succeeded'],
    ['201', 'Created', 'Resource created successfully'],
    ['400', 'Bad Request', 'Invalid input, duplicate email, weak password, invalid rating'],
    ['401', 'Unauthorized', 'Missing/expired/invalid JWT token'],
    ['403', 'Forbidden', 'Trying to modify resource you don\'t own, non-admin accessing admin routes'],
    ['404', 'Not Found', 'Resource doesn\'t exist (meal, user, restaurant, etc.)'],
    ['422', 'Validation Error', 'Request body doesn\'t match expected schema (Pydantic)'],
    ['429', 'Too Many Requests', 'Rate limit exceeded (configured in middleware)'],
    ['500', 'Server Error', 'Unexpected backend error — check logs'],
])

doc.add_page_break()

# ─── 31. Rate Limiting ───
add_colored_heading('31. Rate Limiting & Security')

doc.add_heading('Rate Limiting (middleware.py)', level=2)
doc.add_paragraph('The app uses a custom RateLimitMiddleware:')
add_table(['Limit', 'Window', 'Action'], [
    ['100 requests', 'Per minute per IP', 'Returns 429 Too Many Requests'],
])

doc.add_heading('Security Headers', level=2)
add_table(['Header', 'Value', 'Purpose'], [
    ['X-Content-Type-Options', 'nosniff', 'Prevent MIME type sniffing'],
    ['X-Frame-Options', 'DENY', 'Prevent clickjacking'],
    ['X-XSS-Protection', '1; mode=block', 'Enable XSS filter'],
    ['Strict-Transport-Security', 'max-age=31536000; includeSubDomains', 'Force HTTPS'],
    ['Content-Security-Policy', "default-src 'self'", 'Restrict resource loading'],
])

doc.add_heading('Request Size Limit', level=2)
doc.add_paragraph('Maximum request body size: 10 MB (configured in RequestSizeLimitMiddleware)')

doc.add_heading('Input Sanitization', level=2)
doc.add_paragraph('All text inputs are sanitized via sanitize_string() which strips HTML tags, script tags, and dangerous characters.')

doc.add_heading('Password Requirements', level=2)
doc.add_paragraph('Minimum 8 characters, at least 1 uppercase letter, 1 digit, and 1 special character (!@#$%^&*). Validated via validate_password_strength().')

doc.add_page_break()

# ─── 32. Credentials ───
add_colored_heading('32. Credentials & Access')

doc.add_heading('Admin Access', level=2)
add_table(['Property', 'Value'], [
    ['Admin Panel URL', '{domain}/api/admin-panel'],
    ['Email', 'admin@bo.com'],
    ['Password', 'BoAdmin2026!'],
    ['2FA', 'Demo code displayed on login screen'],
    ['Role', 'admin'],
])

doc.add_heading('Test User', level=2)
add_table(['Property', 'Value'], [
    ['Email', 'test@bo.com'],
    ['Password', 'Test1234!'],
    ['Role', 'user'],
    ['Note', 'Create via POST /api/auth/register'],
])

# ─── FINAL PAGE ───
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(120)
run = p.add_run('END OF API DOCUMENTATION')
run.font.size = Pt(16)
run.font.color.rgb = GREY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('BO Health & Wellness Platform\nAPI Documentation v1.0.0\nBuilt by Flynaut LLC')
run2.font.size = Pt(11)
run2.font.color.rgb = GREY

# SAVE
output_path = '/app/BO_API_Documentation.docx'
doc.save(output_path)
print(f'API Documentation saved to {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')
